"""python-docx 字体信息读/写/替换的测试用例。

覆盖：
- docx_read_text 返回每个 run 的字体信息（font_name/size/bold/italic/color）。
- docx_add_paragraph 的 color 参数 + 字体参数实际写入文档并可读回。
- docx_add_heading 的字体参数。
- docx_replace_text 保留原字体格式（含颜色）。
- docx_replace_text 用 font_* / color 参数应用新字体。
- include_font_info=false 关闭 run 信息。

运行：
    set PYTHONIOENCODING=utf-8
    set PYTHONPATH=<项目根>
    python -m unittest tests.test_docx_font_info -v
"""

import os
import shutil
import tempfile
import unittest

from docx import Document
from docx.shared import RGBColor

from my_hello_agents.tools.builtin import docx_tools


class FontInfoTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp = tempfile.mkdtemp(prefix="docx_font_test_")
        os.environ["DOCX_WORKSPACE"] = cls.tmp
        docx_tools.WORKSPACE = cls.tmp

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    def _new(self, name="t.docx"):
        fp = os.path.join(self.tmp, name)
        docx_tools.run_tool("docx_create", {"file_path": name})
        return fp, name

    # ---- 读取：docx_read_text 返回 run 级字体信息 ----

    def test_read_text_returns_run_font_info(self):
        fp, name = self._new("read.docx")
        # 直接用 python-docx 构造带格式的 run
        doc = Document(fp)
        p = doc.add_paragraph()
        r = p.add_run("红色加粗")
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        r.font.name = "宋体"
        r.font.size = docx_tools.Pt(14)
        doc.save(fp)

        res = docx_tools.run_tool("docx_read_text", {"file_path": name})
        self.assertTrue(res["ok"], res)
        para = res["paragraphs"][0]
        self.assertEqual(para["text"], "红色加粗")
        self.assertTrue(len(para["runs"]) >= 1)
        run = para["runs"][0]
        self.assertEqual(run["text"], "红色加粗")
        self.assertTrue(run["bold"])
        self.assertEqual(run["color"], "FF0000")
        self.assertEqual(run["font_name"], "宋体")
        self.assertEqual(run["font_size"], 14)

    def test_read_text_include_font_info_false(self):
        fp, name = self._new("nofont.docx")
        doc = Document(fp)
        r = doc.add_paragraph().add_run("纯文本")
        r.bold = True
        doc.save(fp)

        res = docx_tools.run_tool(
            "docx_read_text", {"file_path": name, "include_font_info": False}
        )
        self.assertTrue(res["ok"], res)
        para = res["paragraphs"][0]
        self.assertEqual(para["text"], "纯文本")
        self.assertIsNone(para["runs"])

    # ---- 写入：docx_add_paragraph 带 color + 字体 ----

    def test_add_paragraph_writes_color_and_font(self):
        _, name = self._new("write.docx")
        res = docx_tools.run_tool(
            "docx_add_paragraph",
            {
                "file_path": name,
                "text": "高亮文本",
                "color": "red",
                "bold": True,
                "font_name": "Arial",
                "font_size": 16,
            },
        )
        self.assertTrue(res["ok"], res)

        # 读回校验
        doc = Document(os.path.join(self.tmp, name))
        run = doc.paragraphs[0].runs[0]
        self.assertEqual(run.text, "高亮文本")
        self.assertTrue(run.bold)
        self.assertEqual(run.font.name, "Arial")
        self.assertEqual(run.font.size.pt, 16)
        self.assertEqual(str(run.font.color.rgb).upper(), "FF0000")

    def test_add_paragraph_color_hex(self):
        _, name = self._new("hexp.docx")
        res = docx_tools.run_tool(
            "docx_add_paragraph",
            {"file_path": name, "text": "深红", "color": "C00000"},
        )
        self.assertTrue(res["ok"], res)
        doc = Document(os.path.join(self.tmp, name))
        self.assertEqual(str(doc.paragraphs[0].runs[0].font.color.rgb).upper(), "C00000")

    # ---- 写入：docx_add_heading 带字体 ----

    def test_add_heading_with_font_and_color(self):
        _, name = self._new("head.docx")
        res = docx_tools.run_tool(
            "docx_add_heading",
            {
                "file_path": name,
                "text": "标题",
                "level": 1,
                "color": "blue",
                "bold": True,
                "font_size": 20,
            },
        )
        self.assertTrue(res["ok"], res)
        doc = Document(os.path.join(self.tmp, name))
        run = doc.paragraphs[0].runs[0]
        self.assertEqual(run.text, "标题")
        self.assertTrue(run.bold)
        self.assertEqual(run.font.size.pt, 20)
        self.assertEqual(str(run.font.color.rgb).upper(), "0070C0")
        # 标题样式应保留
        self.assertTrue(doc.paragraphs[0].style.name.startswith("Heading"))

    # ---- 替换：保留原字体格式 ----

    def test_replace_keeps_original_color(self):
        fp, name = self._new("keep.docx")
        doc = Document(fp)
        p = doc.add_paragraph()
        r = p.add_run("旧词待替换")
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        r.bold = True
        r.font.name = "楷体"
        doc.save(fp)

        res = docx_tools.run_tool(
            "docx_replace_text",
            {"file_path": name, "find": "旧词", "replace": "新词", "all": True},
        )
        self.assertTrue(res["ok"], res)
        self.assertEqual(res["replaced"], 1)

        doc = Document(fp)
        run = doc.paragraphs[0].runs[0]
        self.assertEqual(run.text, "新词待替换")
        # 字体属性应保留
        self.assertEqual(str(run.font.color.rgb).upper(), "FF0000")
        self.assertTrue(run.bold)
        self.assertEqual(run.font.name, "楷体")

    def test_replace_applies_new_font(self):
        fp, name = self._new("newfont.docx")
        doc = Document(fp)
        p = doc.add_paragraph()
        r = p.add_run("旧词彩色")
        r.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        doc.save(fp)

        res = docx_tools.run_tool(
            "docx_replace_text",
            {
                "file_path": name,
                "find": "旧词",
                "replace": "新词",
                "color": "FF0000",
                "bold": True,
                "font_name": "宋体",
                "font_size": 18,
            },
        )
        self.assertTrue(res["ok"], res)
        doc = Document(fp)
        run = doc.paragraphs[0].runs[0]
        self.assertEqual(run.text, "新词彩色")
        self.assertEqual(str(run.font.color.rgb).upper(), "FF0000")
        self.assertTrue(run.bold)
        self.assertEqual(run.font.name, "宋体")
        self.assertEqual(run.font.size.pt, 18)

    def test_replace_drop_format(self):
        fp, name = self._new("drop.docx")
        doc = Document(fp)
        r = doc.add_paragraph().add_run("旧词")
        r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        r.bold = True
        doc.save(fp)

        res = docx_tools.run_tool(
            "docx_replace_text",
            {
                "file_path": name, "find": "旧词", "replace": "新词",
                "keep_format": False,
            },
        )
        self.assertTrue(res["ok"], res)
        doc = Document(fp)
        run = doc.paragraphs[0].runs[0]
        self.assertEqual(run.text, "新词")
        # 不保留格式：颜色应被清除（None）
        fc = run.font.color.rgb
        self.assertIsNone(fc)

    # ---- schema 校验 ----

    def test_replace_schema_has_font_params(self):
        schemas = {s["function"]["name"]: s["function"]["parameters"]
                   for s in docx_tools.get_tool_schemas()}
        rp = schemas["docx_replace_text"]["properties"]
        for key in ("color", "font_name", "font_size", "bold", "keep_format"):
            self.assertIn(key, rp)
        ap = schemas["docx_add_paragraph"]["properties"]
        self.assertIn("color", ap)
        ah = schemas["docx_add_heading"]["properties"]
        self.assertIn("color", ah)


if __name__ == "__main__":
    unittest.main(verbosity=2)