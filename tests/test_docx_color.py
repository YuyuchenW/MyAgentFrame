"""python-docx 颜色相关工具的测试用例。

覆盖：
- docx_list_colors：列出文档中出现过的字体颜色及次数。
- docx_read_text_by_color：按预设词 / hex 提取指定颜色文字。
- 端到端"先侦察再提取"工作流。
- 参数校验与非法颜色拒绝。

运行：
    set PYTHONIOENCODING=utf-8
    set DOCX_WORKSPACE=<临时目录>
    python -m unittest tests.test_docx_color -v
"""

import os
import shutil
import tempfile
import unittest

from docx import Document
from docx.shared import RGBColor

from my_hello_agents.tools.builtin import docx_tools


def _set_color(run, r, g, b):
    run.font.color.rgb = RGBColor(r, g, b)


def _build_color_docx(path):
    """构造测试文档：
    段落0: 黑字 + 纯红(FF0000) + 默认色
    段落1: 深红(C00000) 一段
    段落2: 蓝色(0070C0) 一段 + 红色(FF0000) 一段（同段双色）
    表格: 第一格黑字，第二格红字
    """
    doc = Document()
    doc.save(path)
    doc = Document(path)
    p0 = doc.add_paragraph("前导文字")
    r0 = p0.add_run("纯红重点")
    _set_color(r0, 0xFF, 0x00, 0x00)
    p0.add_run("，结束。")

    p1 = doc.add_paragraph()
    r1 = p1.add_run("深红标注")
    _set_color(r1, 0xC0, 0x00, 0x00)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("蓝色说明 ")
    _set_color(r2, 0x00, 0x70, 0xC0)
    r2b = p2.add_run("又一句红色")
    _set_color(r2b, 0xFF, 0x00, 0x00)

    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "表黑"
    cell1 = table.cell(0, 1).paragraphs[0]
    cr = cell1.add_run("表红")
    _set_color(cr, 0xFF, 0x00, 0x00)

    doc.save(path)


class DocxColorToolsTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        # 隔离出一个临时工作区，防止污染真实工作目录
        cls.tmp = tempfile.mkdtemp(prefix="docx_color_test_")
        os.environ["DOCX_WORKSPACE"] = cls.tmp
        # 重置全局 WORKSPACE（模块在 import 时已读一次环境变量）
        docx_tools.WORKSPACE = cls.tmp
        cls.fp = os.path.join(cls.tmp, "color_sample.docx")
        _build_color_docx(cls.fp)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp, ignore_errors=True)

    # ---- docx_list_colors ----

    def test_list_colors_finds_all_distinct_colors(self):
        res = docx_tools.run_tool("docx_list_colors", {"file_path": "color_sample.docx"})
        self.assertTrue(res["ok"], res)
        colors = {c["color"]: c["count"] for c in res["colors"]}
        # FF0000 出现 3 次：段落0、段落2、表格
        self.assertEqual(colors.get("FF0000"), 3)
        # C00000 出现 1 次
        self.assertEqual(colors.get("C00000"), 1)
        # 0070C0 出现 1 次
        self.assertEqual(colors.get("0070C0"), 1)
        # 样例文本非空
        for c in res["colors"]:
            self.assertTrue(c["sample"])

    def test_list_colors_exclude_tables(self):
        res = docx_tools.run_tool(
            "docx_list_colors",
            {"file_path": "color_sample.docx", "include_tables": False},
        )
        colors = {c["color"] for c in res["colors"]}
        self.assertIn("FF0000", colors)
        # 仍然能看到正文里 2 处红色（少一个，因为表格里的 1 处被排除）
        cnt = {c["color"]: c["count"] for c in res["colors"]}
        self.assertEqual(cnt["FF0000"], 2)

    # ---- docx_read_text_by_color ----

    def test_extract_red_by_preset_word(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": "red"},
        )
        self.assertTrue(res["ok"], res)
        # 应同时覆盖正文与表格中的红字
        texts = res["plain_text"]
        self.assertIn("纯红重点", texts)
        self.assertIn("又一句红色", texts)
        self.assertIn("表红", texts)
        self.assertEqual(res["matched_runs"], 3)

    def test_extract_red_by_hex_uppercase(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": "FF0000"},
        )
        self.assertTrue(res["ok"], res)
        self.assertIn("纯红重点", res["plain_text"])

    def test_extract_red_by_hex_with_hash(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": "#ff0000"},
        )
        self.assertTrue(res["ok"], res)
        self.assertEqual(res["matched_runs"], 3)

    def test_extract_darkred_does_not_match_preset_red(self):
        """预设 red=FF0000，不能误命中深红 C00000。"""
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": "red"},
        )
        self.assertNotIn("深红标注", res["plain_text"])

    def test_extract_by_exact_darkred_hex(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": "C00000"},
        )
        self.assertTrue(res["ok"], res)
        self.assertEqual(res["matched_runs"], 1)
        self.assertIn("深红标注", res["plain_text"])

    def test_extract_blue(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": "blue"},
        )
        self.assertTrue(res["ok"], res)
        self.assertEqual(res["matched_runs"], 1)
        self.assertIn("蓝色说明", res["plain_text"])

    def test_extract_with_location_metadata(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": "red"},
        )
        items = res["items"]
        locations = {it["location"] for it in items}
        # 表格里的红字 location 应以 table# 开头
        self.assertTrue(
            any(loc.startswith("table#") for loc in locations),
            f"应包含表格位置，实际: {locations}",
        )
        # 正文红字 location 应以 paragraph# 开头
        self.assertTrue(
            any(loc.startswith("paragraph#") for loc in locations),
            f"应包含正文位置，实际: {locations}",
        )

    def test_exclude_tables_flag(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {
                "file_path": "color_sample.docx",
                "color": "red",
                "include_tables": False,
            },
        )
        self.assertEqual(res["matched_runs"], 2)
        self.assertNotIn("表红", res["plain_text"])
        self.assertIn("纯红重点", res["plain_text"])

    def test_invalid_color_rejected(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": "purple"},
        )
        self.assertFalse(res["ok"])
        self.assertIn("不识别的颜色", res["error"])

    def test_nonexistent_file(self):
        res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "nope.docx", "color": "red"},
        )
        self.assertFalse(res["ok"])
        self.assertIn("不存在", res["error"])

    # ---- 端到端"侦察再提取"工作流 ----

    def test_workflow_list_then_extract(self):
        # 1. 模型先侦察颜色分布
        list_res = docx_tools.run_tool(
            "docx_list_colors", {"file_path": "color_sample.docx"}
        )
        self.assertTrue(list_res["ok"])
        # 找到出现次数最多的颜色 hex（模拟模型据此决策）
        top = max(list_res["colors"], key=lambda c: c["count"])
        self.assertEqual(top["color"], "FF0000")
        # 2. 再按该 hex 精确提取
        extract_res = docx_tools.run_tool(
            "docx_read_text_by_color",
            {"file_path": "color_sample.docx", "color": top["color"]},
        )
        self.assertTrue(extract_res["ok"])
        self.assertEqual(extract_res["matched_runs"], 3)
        self.assertIn("纯红重点", extract_res["plain_text"])

    # ---- OpenAI function-calling schema 校验 ----

    def test_tools_registered_with_valid_schema(self):
        schemas = docx_tools.get_tool_schemas()
        names = {
            s["function"]["name"] for s in schemas
        }
        self.assertIn("docx_list_colors", names)
        self.assertIn("docx_read_text_by_color", names)
        for s in schemas:
            self.assertEqual(s["type"], "function")
            self.assertIn("parameters", s["function"])


if __name__ == "__main__":
    unittest.main(verbosity=2)