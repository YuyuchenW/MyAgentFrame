"""python-docx 工具集 —— 让大模型通过 Function Calling 操作 Word 文档。

设计要点：
- 细粒度多工具：每个动作一个独立工具，schema 清晰。
- 一次性文件制（无状态）：每个工具接收 file_path，内部打开、修改、保存。
- 全功能：覆盖段落、标题、表格、图片、页眉页脚、页边距、查找替换等。
- 安全：工作区根目录限制、覆盖保护、读取截断、规模上限。
"""

import os
import json
from typing import Any, Callable, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ==================== 全局配置 / 安全限制 ====================

# 工作区根目录：所有 file_path/image_path 都被限定在该目录内。
# 可通过环境变量 DOCX_WORKSPACE 覆盖，默认当前工作目录。
WORKSPACE = os.path.abspath(os.getenv("DOCX_WORKSPACE", os.getcwd()))

# 规模上限，防止 LLM 死循环塞爆文档
MAX_PARAGRAPHS = 50000
MAX_TABLE_ROWS = 5000
MAX_TABLE_COLS = 500
MAX_READ_CHARS = 80000  # 读取文本时单次返回字符上限

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ==================== 路径安全校验 ====================

def _safe_path(p: str, must_exist: bool = False) -> str:
    """将相对路径解析到 WORKSPACE 内，禁止绝对路径与 .. 越界。"""
    if not p:
        raise ValueError("路径不能为空")

    # 不允许绝对路径，统一相对 WORKSPACE
    full = os.path.abspath(os.path.join(WORKSPACE, p))
    if os.path.isabs(p):
        # 允许绝对路径，但必须在 WORKSPACE 之下
        full = os.path.abspath(p)

    if not full.startswith(WORKSPACE):
        raise PermissionError(
            f"路径越界，只允许在 workspace 内操作: {WORKSPACE}"
        )

    if must_exist and not os.path.exists(full):
        raise FileNotFoundError(f"文件不存在: {p}")

    return full


def _check_size(doc: "Document") -> None:
    if len(doc.paragraphs) > MAX_PARAGRAPHS:
        raise ValueError(f"段落数超上限 {MAX_PARAGRAPHS}")


# ==================== 字体信息统一处理 ====================

def _run_font_info(run) -> Dict[str, Any]:
    """提取一个 run 的字体信息为可序列化 dict（供模型理解文档格式）。"""
    fc = run.font.color
    color: Optional[str] = None
    if fc and fc.type is not None and getattr(fc, "rgb", None) is not None:
        color = str(fc.rgb).upper()
    elif fc and getattr(fc, "theme_color", None) is not None:
        color = f"theme:{fc.theme_color}"
    size = None
    if run.font.size is not None:
        size = run.font.size.pt
    return {
        "text": run.text,
        "font_name": run.font.name,
        "font_size": size,
        "bold": bool(run.bold) if run.bold is not None else None,
        "italic": bool(run.italic) if run.italic is not None else None,
        "underline": bool(run.underline) if run.underline is not None else None,
        "color": color,
    }


def _apply_font(
    run,
    *,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    color: Optional[str] = None,
) -> None:
    """把可选字体参数应用到 run。color 接受预设词或 hex。"""
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if color:
        rgb = _color_label_to_hex(color)
        if rgb is None:
            raise ValueError(f"不识别的颜色: {color}，可用 red/green/blue 或 hex 如 FF0000")
        run.font.color.rgb = RGBColor.from_string(rgb)


def _clear_font(run) -> None:
    """清除 run 的显式字体属性，使其回落到样式默认。"""
    run.bold = None
    run.italic = None
    run.underline = None
    run.font.name = None
    run.font.size = None
    # 清除显式颜色：删除 run 的 rPr 下的 color 子元素
    rpr = run._element.rPr
    if rpr is not None:
        from docx.oxml.ns import qn
        for tag in ("w:color",):
            el = rpr.find(qn(tag))
            if el is not None:
                rpr.remove(el)


# ==================== 工具注册表 ====================

DOCX_TOOLS: Dict[str, Dict[str, Any]] = {}


def _register(name: str, description: str, parameters: Dict[str, Any]) -> Callable:
    """装饰器：把函数注册为 OpenAI function-calling 工具。"""

    def deco(fn: Callable) -> Callable:
        DOCX_TOOLS[name] = {
            "schema": {
                "type": "function",
                "function": {
                    "name": name,
                    "description": description,
                    "parameters": parameters,
                },
            },
            "run": fn,
        }
        return fn

    return deco


# ==================== 写入类工具（11 个） ====================

@_register(
    "docx_create",
    "新建一个空的 Word 文档。若目标文件已存在，需显式 overwrite=true 才会覆盖。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "相对 workspace 的文档路径"},
            "overwrite": {"type": "boolean", "description": "已存在时是否覆盖，默认 false"},
        },
        "required": ["file_path"],
    },
)
def docx_create(file_path: str, overwrite: bool = False) -> Dict[str, Any]:
    full = _safe_path(file_path)
    if os.path.exists(full) and not overwrite:
        raise FileExistsError(f"文件已存在: {file_path}，需 overwrite=true")
    os.makedirs(os.path.dirname(full), exist_ok=True)
    Document().save(full)
    return {"ok": True, "file_path": file_path}


@_register(
    "docx_add_heading",
    "向已有文档追加一个标题段落，自动应用 Heading 样式，并支持设置标题字体格式。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "text": {"type": "string"},
            "level": {"type": "integer", "minimum": 1, "maximum": 9, "description": "标题级别，1=一级标题"},
            "bold": {"type": "boolean"},
            "italic": {"type": "boolean"},
            "underline": {"type": "boolean"},
            "font_name": {"type": "string"},
            "font_size": {"type": "number", "description": "字号（磅）"},
            "color": {"type": "string", "description": "颜色，预设词如 red 或 hex 如 FF0000"},
        },
        "required": ["file_path", "text", "level"],
    },
)
def docx_add_heading(
    file_path: str,
    text: str,
    level: int = 1,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    color: Optional[str] = None,
) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _apply_font(
            run, bold=bold, italic=italic, underline=underline,
            font_name=font_name, font_size=font_size, color=color,
        )
    _check_size(doc)
    doc.save(full)
    return {"ok": True, "paragraphs": len(doc.paragraphs)}


@_register(
    "docx_add_paragraph",
    "向已有文档追加一个段落，支持 run 级格式（加粗/斜体/下划线/字体/字号/颜色）与段对齐、样式。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "text": {"type": "string"},
            "bold": {"type": "boolean"},
            "italic": {"type": "boolean"},
            "underline": {"type": "boolean"},
            "font_name": {"type": "string", "description": "如 '宋体'、'Arial'"},
            "font_size": {"type": "number", "description": "字号（磅）"},
            "color": {"type": "string", "description": "颜色，预设词如 red/green/blue 或 hex 如 FF0000"},
            "style": {"type": "string", "description": "段落样式名，如 'List Bullet'、'Title'"},
            "align": {"type": "string", "enum": ["left", "center", "right", "justify"]},
        },
        "required": ["file_path", "text"],
    },
)
def docx_add_paragraph(
    file_path: str,
    text: str,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    color: Optional[str] = None,
    style: Optional[str] = None,
    align: Optional[str] = None,
) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    _apply_font(
        run, bold=bold, italic=italic, underline=underline,
        font_name=font_name, font_size=font_size, color=color,
    )
    if align:
        p.alignment = ALIGN_MAP[align]
    _check_size(doc)
    doc.save(full)
    return {"ok": True, "paragraphs": len(doc.paragraphs)}


@_register(
    "docx_add_table",
    "向已有文档追加一个表格。rows 为二维字符串数组，第一项可作为表头。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "rows": {
                "type": "array",
                "items": {"type": "array", "items": {"type": "string"}},
                "description": "二维字符串数组，每个子数组是一行",
            },
            "header": {"type": "boolean", "description": "第一行是否作为表头（加粗），默认 true"},
            "style": {"type": "string", "description": "表格样式名，如 'Table Grid'"},
        },
        "required": ["file_path", "rows"],
    },
)
def docx_add_table(
    file_path: str,
    rows: List[List[str]],
    header: bool = True,
    style: str = "Table Grid",
) -> Dict[str, Any]:
    if not rows:
        raise ValueError("rows 不能为空")
    n_rows, n_cols = len(rows), max(len(r) for r in rows)
    if n_rows > MAX_TABLE_ROWS or n_cols > MAX_TABLE_COLS:
        raise ValueError(f"表格规模超限 (max {MAX_TABLE_ROWS}x{MAX_TABLE_COLS})")

    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        table.style = style
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j in range(len(row)):
            cell = table.cell(i, j)
            cell.text = row[j]
            if header and i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.save(full)
    return {"ok": True, "tables": len(doc.tables), "rows": n_rows, "cols": n_cols}


@_register(
    "docx_set_table_cell",
    "修改已有文档中指定表格的指定单元格内容。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "table_index": {"type": "integer", "minimum": 0, "description": "第几个表格，从 0 开始"},
            "row": {"type": "integer", "minimum": 0},
            "col": {"type": "integer", "minimum": 0},
            "text": {"type": "string"},
        },
        "required": ["file_path", "table_index", "row", "col", "text"],
    },
)
def docx_set_table_cell(
    file_path: str, table_index: int, row: int, col: int, text: str
) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    if table_index >= len(doc.tables):
        raise IndexError(f"无此表格: table_index={table_index}")
    doc.tables[table_index].cell(row, col).text = text
    doc.save(full)
    return {"ok": True}


@_register(
    "docx_add_image",
    "向已有文档插入一张本地图片。image_path 也必须在 workspace 内。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "image_path": {"type": "string", "description": "workspace 内的图片路径"},
            "width_inches": {"type": "number", "description": "显示宽度（英寸），不传则按原图"},
            "height_inches": {"type": "number", "description": "显示高度（英寸）"},
        },
        "required": ["file_path", "image_path"],
    },
)
def docx_add_image(
    file_path: str,
    image_path: str,
    width_inches: Optional[float] = None,
    height_inches: Optional[float] = None,
) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    img_full = _safe_path(image_path, must_exist=True)
    doc = Document(full)
    kwargs: Dict[str, Any] = {}
    if width_inches:
        kwargs["width"] = Inches(width_inches)
    if height_inches:
        kwargs["height"] = Inches(height_inches)
    doc.add_picture(img_full, **kwargs)
    doc.save(full)
    return {"ok": True}


@_register(
    "docx_add_page_break",
    "向已有文档插入分页符。",
    {
        "type": "object",
        "properties": {"file_path": {"type": "string"}},
        "required": ["file_path"],
    },
)
def docx_add_page_break(file_path: str) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    doc.add_page_break()
    doc.save(full)
    return {"ok": True}


@_register(
    "docx_set_margins",
    "设置文档所有节的页边距（单位 cm）。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "top": {"type": "number"},
            "bottom": {"type": "number"},
            "left": {"type": "number"},
            "right": {"type": "number"},
        },
        "required": ["file_path"],
    },
)
def docx_set_margins(
    file_path: str,
    top: Optional[float] = None,
    bottom: Optional[float] = None,
    left: Optional[float] = None,
    right: Optional[float] = None,
) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    for section in doc.sections:
        if top is not None:
            section.top_margin = Cm(top)
        if bottom is not None:
            section.bottom_margin = Cm(bottom)
        if left is not None:
            section.left_margin = Cm(left)
        if right is not None:
            section.right_margin = Cm(right)
    doc.save(full)
    return {"ok": True, "sections": len(doc.sections)}


@_register(
    "docx_add_header_footer",
    "为文档指定节添加页眉/页脚文本。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "section": {"type": "integer", "minimum": 0, "description": "节索引，默认 0"},
            "header_text": {"type": "string"},
            "footer_text": {"type": "string"},
        },
        "required": ["file_path"],
    },
)
def docx_add_header_footer(
    file_path: str,
    section: int = 0,
    header_text: Optional[str] = None,
    footer_text: Optional[str] = None,
) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    if section >= len(doc.sections):
        raise IndexError(f"无此节: section={section}")
    sec = doc.sections[section]
    if header_text is not None:
        sec.header.paragraphs[0].text = header_text
    if footer_text is not None:
        sec.footer.paragraphs[0].text = footer_text
    doc.save(full)
    return {"ok": True}


@_register(
    "docx_delete_paragraph",
    "按索引删除文档中的段落。索引从 0 开始。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "index": {"type": "integer", "minimum": 0},
        },
        "required": ["file_path", "index"],
    },
)
def docx_delete_paragraph(file_path: str, index: int) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    if index >= len(doc.paragraphs):
        raise IndexError(f"无此段落: index={index}")
    p = doc.paragraphs[index]
    p._element.getparent().remove(p._element)
    doc.save(full)
    return {"ok": True, "paragraphs": len(doc.paragraphs)}


@_register(
    "docx_replace_text",
    "在文档中进行查找替换，默认保留每个 run 原有字体格式。"
    "可指定 keep_format=false 丢弃原格式，也可通过 font_* / color 参数"
    "为替换后的文字应用新的字体格式。默认全量替换，all=false 只替换第一个匹配。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "find": {"type": "string"},
            "replace": {"type": "string"},
            "all": {"type": "boolean", "description": "是否替换全部，默认 true"},
            "keep_format": {"type": "boolean", "description": "是否保留命中原 run 的字体格式，默认 true"},
            "font_name": {"type": "string", "description": "为替换后文字设置字体名"},
            "font_size": {"type": "number", "description": "为替换后文字设置字号（磅）"},
            "bold": {"type": "boolean"},
            "italic": {"type": "boolean"},
            "underline": {"type": "boolean"},
            "color": {"type": "string", "description": "为替换后文字设置颜色，如 red 或 FF0000"},
        },
        "required": ["file_path", "find", "replace"],
    },
)
def docx_replace_text(
    file_path: str,
    find: str,
    replace: str,
    all: bool = True,
    keep_format: bool = True,
    font_name: Optional[str] = None,
    font_size: Optional[float] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    color: Optional[str] = None,
) -> Dict[str, Any]:
    if not find:
        raise ValueError("find 不能为空")
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    count = 0

    has_new_font = any(
        v is not None
        for v in (font_name, font_size, bold, italic, underline, color)
    )

    def _apply_new_font(run) -> None:
        if not has_new_font:
            return
        _apply_font(
            run, bold=bold, italic=italic, underline=underline,
            font_name=font_name, font_size=font_size, color=color,
        )

    def _replace_in_paragraph(p) -> int:
        nonlocal count
        if find not in p.text:
            return 0
        runs = p.runs
        if not runs:
            return 0
        # 简化策略：在段落文本层面做替换；命中所在的 run 保留其字体属性，
        # 把替换后的整段文本写入该 run，其余 run 清空文本但不改其字体属性。
        # 注意：python-docx 的 p.runs 每次调用返回新包装对象，对象 identity 不稳定，
        # 因此用下标定位命中 run，清空时按下标比较，绝不能用 `is`。
        hit_idx = -1
        run_font = None
        for i, r in enumerate(runs):
            if r.text and find in r.text:
                hit_idx = i
                if keep_format:
                    run_font = _run_font_info(r)
                break
        if hit_idx < 0:
            # 跨 run 命中：取第一个非空 run 的字体作为模板
            for i, r in enumerate(runs):
                if r.text:
                    hit_idx = i
                    if keep_format:
                        run_font = _run_font_info(r)
                    break
            if hit_idx < 0:
                return 0
        times = p.text.count(find) if all else 1
        new_text = p.text.replace(find, replace) if all else p.text.replace(find, replace, 1)
        target = runs[hit_idx]
        target.text = new_text
        # 清空其余 run 的文本（字体属性保留，不清除）
        for i, r in enumerate(runs):
            if i == hit_idx:
                continue
            r.text = ""
        # 保留原字体：python-docx 修改 run.text 不会清除 run.font 属性，
        # 因此 target 的字体属性天然保留；这里显式重设以稳定跨版本行为。
        if keep_format and run_font:
            if run_font.get("font_name"):
                target.font.name = run_font["font_name"]
            if run_font.get("font_size") is not None:
                target.font.size = Pt(run_font["font_size"])
            if run_font.get("bold") is not None:
                target.bold = run_font["bold"]
            if run_font.get("italic") is not None:
                target.italic = run_font["italic"]
            if run_font.get("underline") is not None:
                target.underline = run_font["underline"]
            if run_font.get("color"):
                c = run_font["color"]
                if not c.startswith("theme:"):
                    target.font.color.rgb = RGBColor.from_string(c)
        else:
            # 不保留格式：清除命中 run 的显式字体，使其回落到段落样式默认
            _clear_font(target)
        _apply_new_font(target)
        count += times
        return times

    # 正文段落
    for p in doc.paragraphs:
        _replace_in_paragraph(p)
    # 表格单元格中的段落
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p)
    # 页眉页脚
    for section in doc.sections:
        for p in list(section.header.paragraphs) + list(section.footer.paragraphs):
            _replace_in_paragraph(p)
    doc.save(full)
    return {"ok": True, "replaced": count}


# ==================== 读取类工具（4 个，结果需截断） ====================

@_register(
    "docx_read_text",
    "读取已有 Word 文档的全文文本，按段落返回，并附带每个 run 的字体信息"
    "（字体名/字号/加粗/斜体/下划线/颜色）。超长会自动截断并提示。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "include_font_info": {
                "type": "boolean",
                "description": "是否返回每个 run 的字体信息，默认 true",
            },
        },
        "required": ["file_path"],
    },
)
def docx_read_text(file_path: str, include_font_info: bool = True) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    paragraphs: List[Dict[str, Any]] = []
    text_parts: List[str] = []
    total_len = 0
    truncated = False

    for i, p in enumerate(doc.paragraphs):
        if total_len > MAX_READ_CHARS:
            truncated = True
            break
        runs: List[Dict[str, Any]] = []
        for run in p.runs:
            info = _run_font_info(run)
            total_len += len(info["text"])
            if include_font_info:
                runs.append(info)
        style = (p.style.name if p.style else None) or None
        paragraphs.append({
            "index": i,
            "text": p.text,
            "style": style,
            "runs": runs if include_font_info else None,
        })
        text_parts.append(p.text)

    text = "\n".join(text_parts)
    if len(text) > MAX_READ_CHARS:
        text = text[:MAX_READ_CHARS]
        truncated = True
    return {
        "ok": True,
        "paragraphs": paragraphs,
        "text": text,
        "truncated": truncated,
        "total_paragraphs": len(doc.paragraphs),
        "hint": "内容过长已截断，可用 docx_list_structure 查看大纲后再分段处理"
        if truncated else None,
    }


@_register(
    "docx_read_tables",
    "读取文档中所有表格，返回 JSON 数组。",
    {
        "type": "object",
        "properties": {"file_path": {"type": "string"}},
        "required": ["file_path"],
    },
)
def docx_read_tables(file_path: str) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    tables = []
    for t in doc.tables:
        tables.append([[cell.text for cell in row.cells] for row in t.rows])
    return {"tables": tables, "count": len(tables)}


@_register(
    "docx_list_structure",
    "列出文档大纲（标题层级与段落数），用于快速了解文档结构。",
    {
        "type": "object",
        "properties": {"file_path": {"type": "string"}},
        "required": ["file_path"],
    },
)
def docx_list_structure(file_path: str) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    outline: List[Dict[str, Any]] = []
    for i, p in enumerate(doc.paragraphs):
        style = (p.style.name or "") if p.style else ""
        if style.startswith("Heading"):
            try:
                level = int(style.replace("Heading", "").strip() or "1")
            except ValueError:
                level = 0
            outline.append({"index": i, "level": level, "text": p.text})
    return {
        "outline": outline,
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
    }


@_register(
    "docx_read_metadata",
    "读取文档元数据：作者、标题、段落数、表格数等。",
    {
        "type": "object",
        "properties": {"file_path": {"type": "string"}},
        "required": ["file_path"],
    },
)
def docx_read_metadata(file_path: str) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    cp = doc.core_properties
    return {
        "author": cp.author,
        "title": cp.title,
        "subject": cp.subject,
        "created": cp.created.isoformat() if cp.created else None,
        "modified": cp.modified.isoformat() if cp.modified else None,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
    }


# 颜色预设：覆盖 Word 常用"红字"等口语化表述，统一映射为 RGB hex（大写）。
COLOR_PRESETS: Dict[str, str] = {
    "red": "FF0000",
    "darkred": "8B0000",
    "green": "00B050",
    "darkgreen": "006400",
    "blue": "0070C0",
    "darkblue": "0000FF",
    "yellow": "FFFF00",
    "orange": "ED7D31",
    "black": "000000",
    "gray": "808080",
    "white": "FFFFFF",
}


def _color_label_to_hex(color: str) -> Optional[str]:
    """把预设词或 '#RRGGBB'/'RRGGBB' 统一为大写 6 位 hex，无法识别返回 None。"""
    c = color.strip().lstrip("#").upper()
    if c in {v.upper() for v in COLOR_PRESETS.values()}:
        return c
    preset = COLOR_PRESETS.get(color.lower())
    if preset:
        return preset
    if len(c) == 6 and all(ch in "0123456789ABCDEF" for ch in c):
        return c
    return None


@_register(
    "docx_list_colors",
    "列出文档中出现过的所有字体颜色及其出现次数与样例文本，"
    "用于在未知文档中先侦察有哪些颜色再决定提取哪种颜色（如找红字）。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "include_tables": {"type": "boolean", "description": "是否扫描表格，默认 true"},
            "include_header_footer": {"type": "boolean", "description": "是否扫描页眉页脚，默认 true"},
        },
        "required": ["file_path"],
    },
)
def docx_list_colors(
    file_path: str,
    include_tables: bool = True,
    include_header_footer: bool = True,
) -> Dict[str, Any]:
    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    color_map: Dict[str, Dict[str, Any]] = {}

    def _scan(p):
        for run in p.runs:
            fc = run.font.color
            rgb = None
            if fc and fc.type is not None and getattr(fc, "rgb", None) is not None:
                rgb = str(fc.rgb).upper()
            elif fc and getattr(fc, "theme_color", None) is not None:
                rgb = f"theme:{fc.theme_color}"
            if rgb is None or not (run.text or "").strip():
                continue
            entry = color_map.setdefault(rgb, {"color": rgb, "count": 0, "sample": ""})
            entry["count"] += 1
            if not entry["sample"]:
                entry["sample"] = run.text[:50]

    for p in doc.paragraphs:
        _scan(p)
    if include_tables:
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _scan(p)
    if include_header_footer:
        for section in doc.sections:
            for p in section.header.paragraphs:
                _scan(p)
            for p in section.footer.paragraphs:
                _scan(p)

    items = sorted(color_map.values(), key=lambda x: -x["count"])
    return {"ok": True, "colors": items, "total": len(items)}


@_register(
    "docx_read_text_by_color",
    "提取文档中字体颜色匹配指定颜色的文字，常用于读取'红字'等高亮内容。"
    "color 可用预设词（red/green/blue/yellow/orange/black/gray/white 等）"
    "或 6 位 hex（如 FF0000、#C00000）。返回所有匹配 run 的文本及其位置。",
    {
        "type": "object",
        "properties": {
            "file_path": {"type": "string"},
            "color": {
                "type": "string",
                "description": "目标颜色，预设词（如 red）或 hex（如 FF0000、#C00000）",
            },
            "include_tables": {"type": "boolean", "description": "是否扫描表格内文字，默认 true"},
            "include_header_footer": {"type": "boolean", "description": "是否扫描页眉页脚，默认 true"},
        },
        "required": ["file_path", "color"],
    },
)
def docx_read_text_by_color(
    file_path: str,
    color: str,
    include_tables: bool = True,
    include_header_footer: bool = True,
) -> Dict[str, Any]:
    target = _color_label_to_hex(color)
    if target is None:
        raise ValueError(f"不识别的颜色: {color}，可用 red/green/blue 或 hex 如 FF0000")

    full = _safe_path(file_path, must_exist=True)
    doc = Document(full)
    matches: List[Dict[str, Any]] = []
    total_len = 0

    def _scan_paragraph(p, location: str) -> None:
        nonlocal total_len
        for run in p.runs:
            if total_len > MAX_READ_CHARS:
                return
            fc = run.font.color
            rgb = None
            if fc and fc.type is not None and getattr(fc, "rgb", None) is not None:
                rgb = str(fc.rgb).upper()
            elif fc and getattr(fc, "theme_color", None) is not None:
                rgb = f"theme:{fc.theme_color}"
            if rgb is None:
                continue
            if rgb == target or (target in COLOR_PRESETS.values() and rgb == target):
                text = run.text or ""
                if text:
                    matches.append({"location": location, "color": rgb, "text": text})
                    total_len += len(text)

    # 正文段落
    for i, p in enumerate(doc.paragraphs):
        _scan_paragraph(p, f"paragraph#{i}")
    # 表格
    if include_tables:
        for ti, t in enumerate(doc.tables):
            for ri, row in enumerate(t.rows):
                for ci, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        _scan_paragraph(p, f"table#{ti}:{ri},{ci}")
    # 页眉页脚
    if include_header_footer:
        for si, section in enumerate(doc.sections):
            for p in section.header.paragraphs:
                _scan_paragraph(p, f"header#section{si}")
            for p in section.footer.paragraphs:
                _scan_paragraph(p, f"footer#section{si}")

    return {
        "ok": True,
        "color": color,
        "matched_runs": len(matches),
        "items": matches,
        "plain_text": "".join(m["text"] for m in matches),
        "truncated": total_len > MAX_READ_CHARS,
    }


# ==================== 对外辅助 ====================

def get_tool_schemas() -> List[Dict[str, Any]]:
    """返回 OpenAI function-calling 格式的工具 schema 列表。"""
    return [v["schema"] for v in DOCX_TOOLS.values()]

# TODO 后续删除
def register_docx_tools(registry) -> None:
    """Register all docx functions in a ToolRegistry."""
    for item in DOCX_TOOLS.values():
        function_schema = item["schema"]["function"]
        registry.register_schema_function(
            name=function_schema["name"],
            description=function_schema["description"],
            parameters=function_schema["parameters"],
            func=item["run"],
        )


def create_docx_tool_registry(circuit_breaker=None):
    """Create a ToolRegistry preloaded with all docx tools."""
    from my_hello_agents.tools.registry import ToolRegistry

    registry = ToolRegistry(circuit_breaker=circuit_breaker)
    register_docx_tools(registry)
    return registry


def run_tool(name: str, arguments: Dict[str, Any]) -> Dict[str, Any]:
    """按名字执行某个工具，返回结果 dict（出错时带 ok=False/error）。"""
    tool = DOCX_TOOLS.get(name)
    if not tool:
        return {"ok": False, "error": f"未知工具: {name}"}
    try:
        return tool["run"](**arguments) or {"ok": True}
    except Exception as e:
        return {"ok": False, "error": f"{type(e).__name__}: {e}"}


if __name__ == "__main__":
    # 自检：列出所有注册工具
    print(f"workspace: {WORKSPACE}")
    print(f"已注册 {len(DOCX_TOOLS)} 个 docx 工具:")
    for n in DOCX_TOOLS:
        print(" -", n)
